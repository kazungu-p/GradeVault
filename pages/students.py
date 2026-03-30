import customtkinter as ctk
import os
from tkinter import filedialog
from utils.theme import *
from utils.pdf_classlist import generate_class_list
from utils.importer import read_students_from_file, sample_csv_template
from routes.students import (
    get_students, get_classes, create_student,
    update_student, transfer_student, archive_student,
)

PAGE_SIZE = 50  # students per page


class StudentsPage(ctk.CTkFrame):
    def __init__(self, parent):
        super().__init__(parent, fg_color=BG)
        self._students    = []
        self._classes     = []
        self._page        = 0
        self._show_stream_col = False
        self._thead       = None
        self._build()
        self._load()

    # ── Layout ───────────────────────────────────────────────
    def _build(self):
        self.pack(fill="both", expand=True, padx=20, pady=20)

        # Header
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", pady=(0, 10))
        heading(header, "Students").pack(side="left")
        primary_btn(header, "+ Add student",
                    command=self._open_add_form,
                    width=130).pack(side="right")
        ghost_btn(header, "Print / export",
                  command=self._print_class_list,
                  width=120).pack(side="right", padx=(0, 8))
        ghost_btn(header, "Import CSV/Excel",
                  command=self._open_import,
                  width=130).pack(side="right", padx=(0, 8))

        # Filter row
        filters = ctk.CTkFrame(self, fg_color="transparent")
        filters.pack(fill="x", pady=(0, 8))

        # Search
        self._search_var = ctk.StringVar()
        self._search_var.trace_add("write", lambda *_: self._reset_load())
        ctk.CTkEntry(filters, placeholder_text="Search name or adm. no.",
                     textvariable=self._search_var, width=220,
                     border_color=BORDER, fg_color=SURFACE
                     ).pack(side="left", padx=(0, 8))

        # Class filter
        self._class_var = ctk.StringVar(value="All classes")
        self._class_menu = ctk.CTkOptionMenu(
            filters, variable=self._class_var,
            values=["All classes"],
            command=self._on_class_filter,
            width=150, fg_color=SURFACE,
            button_color=BORDER, button_hover_color=BORDER,
            text_color=TEXT, dropdown_fg_color=SURFACE)
        self._class_menu.pack(side="left", padx=(0, 8))

        # Stream filter (hidden until class selected with streams)
        self._stream_var = ctk.StringVar(value="All streams")
        self._stream_menu = ctk.CTkOptionMenu(
            filters, variable=self._stream_var,
            values=["All streams"],
            command=lambda _: self._reset_load(),
            width=130, fg_color=SURFACE,
            button_color=BORDER, button_hover_color=BORDER,
            text_color=TEXT, dropdown_fg_color=SURFACE)
        # Pack later only if needed

        # Gender filter
        self._gender_var = ctk.StringVar(value="All")
        ctk.CTkOptionMenu(
            filters, variable=self._gender_var,
            values=["All", "Male", "Female"],
            command=lambda _: self._reset_load(),
            width=90, fg_color=SURFACE,
            button_color=BORDER, button_hover_color=BORDER,
            text_color=TEXT, dropdown_fg_color=SURFACE
        ).pack(side="left", padx=(0, 8))

        # Status filter
        self._status_var = ctk.StringVar(value="Active")
        ctk.CTkOptionMenu(
            filters, variable=self._status_var,
            values=["Active", "Archived"],
            command=lambda _: self._reset_load(),
            width=100, fg_color=SURFACE,
            button_color=BORDER, button_hover_color=BORDER,
            text_color=TEXT, dropdown_fg_color=SURFACE
        ).pack(side="left", padx=(0, 8))

        self._count_label = muted(filters, "")
        self._count_label.pack(side="right")

        # Table card
        tcard = card(self)
        tcard.pack(fill="both", expand=True)

        thead = ctk.CTkFrame(tcard, fg_color="#F3F4F6", corner_radius=0)
        thead.pack(fill="x", padx=1, pady=(1, 0))
        self._thead = thead
        self._render_thead()

        self._body = ctk.CTkScrollableFrame(
            tcard, fg_color=SURFACE, corner_radius=0)
        self._body.pack(fill="both", expand=True, padx=1, pady=(0, 1))

        # Pagination bar
        self._pager = ctk.CTkFrame(tcard, fg_color="transparent", height=36)
        self._pager.pack(fill="x", padx=8, pady=4)

    def _render_thead(self):
        if not self._thead:
            return
        for w in self._thead.winfo_children():
            w.destroy()
        cols = [("Admission no.", 130), ("Full name", 230), ("Class", 100)]
        if self._show_stream_col:
            cols.append(("Stream", 80))
        cols += [("Gender", 60), ("Actions", 180)]
        for col_name, col_w in cols:
            ctk.CTkLabel(self._thead, text=col_name,
                         font=("", 11, "bold"), text_color=TEXT_MUTED,
                         width=col_w, anchor="w"
                         ).pack(side="left", padx=(12, 0), pady=8)

    # ── Filters ──────────────────────────────────────────────
    def _on_class_filter(self, _=None):
        self._update_stream_menu()
        self._reset_load()

    def _update_stream_menu(self):
        sel = self._class_var.get()
        if sel == "All classes":
            self._stream_menu.pack_forget()
            return
        # Find streams for selected class name
        cls_name = sel
        streams = sorted(set(
            c["stream"] for c in self._classes
            if c["name"] == cls_name and c.get("stream")
        ))
        if len(streams) > 1:
            self._stream_var.set("All streams")
            self._stream_menu.configure(values=["All streams"] + streams)
            self._stream_menu.pack(side="left", padx=(0, 8),
                                   after=self._class_menu)
        else:
            self._stream_menu.pack_forget()

    def _reset_load(self):
        self._page = 0
        self._load()

    # ── Data ─────────────────────────────────────────────────
    def _load(self):
        search = self._search_var.get().strip()
        status = "active" if self._status_var.get() == "Active" else "archived"

        # Resolve class filter → class_id
        class_id = None
        sel_class = self._class_var.get()
        if sel_class != "All classes":
            sel_stream = self._stream_var.get()
            if sel_stream and sel_stream != "All streams":
                match = next((c for c in self._classes
                              if c["name"] == sel_class
                              and c.get("stream") == sel_stream), None)
            else:
                match = next((c for c in self._classes
                              if c["name"] == sel_class), None)
            if match:
                # If "All streams" selected, pass class name not id
                if sel_stream == "All streams":
                    class_id = None  # handled below
                else:
                    class_id = match["id"]

        # Get all students for the selected class name (any stream)
        all_students = get_students(
            search=search, class_id=class_id, status=status)

        # If "all streams" for a class name, filter by name
        if sel_class != "All classes" and self._stream_var.get() == "All streams":
            all_students = [s for s in get_students(search=search, status=status)
                            if s.get("class_name") == sel_class]

        # Gender filter
        gender_sel = self._gender_var.get()
        if gender_sel == "Male":
            all_students = [s for s in all_students if s.get("gender") == "M"]
        elif gender_sel == "Female":
            all_students = [s for s in all_students if s.get("gender") == "F"]

        self._all_students = all_students

        # Refresh class dropdown
        self._classes = get_classes()
        unique_names = []
        for c in self._classes:
            if c["name"] not in unique_names:
                unique_names.append(c["name"])
        self._class_menu.configure(
            values=["All classes"] + unique_names)

        # Update stream column visibility
        streams = set(s.get("stream") for s in all_students if s.get("stream"))
        self._show_stream_col = len(streams) > 1
        self._render_thead()

        total = len(all_students)
        pages = max(1, (total + PAGE_SIZE - 1) // PAGE_SIZE)
        self._page = min(self._page, pages - 1)
        start = self._page * PAGE_SIZE
        self._students = all_students[start:start + PAGE_SIZE]

        self._count_label.configure(
            text=f"{total} student(s)  —  page {self._page+1} of {pages}")
        self._render_rows()
        self._render_pager(pages)

    def _render_pager(self, pages):
        for w in self._pager.winfo_children():
            w.destroy()
        if pages <= 1:
            return

        def go(p):
            self._page = p
            self._load()

        ghost_btn(self._pager, "← Prev",
                  command=lambda: go(max(0, self._page - 1)),
                  width=80).pack(side="left", padx=(0, 6))

        for p in range(pages):
            is_cur = p == self._page
            ctk.CTkButton(
                self._pager, text=str(p + 1), width=32, height=28,
                fg_color=ACCENT if is_cur else "transparent",
                text_color="white" if is_cur else TEXT_MUTED,
                hover_color=ACCENT_BG, corner_radius=6, font=("", 12),
                command=lambda pg=p: go(pg),
            ).pack(side="left", padx=2)

        ghost_btn(self._pager, "Next →",
                  command=lambda: go(min(pages - 1, self._page + 1)),
                  width=80).pack(side="left", padx=(6, 0))

    def _render_rows(self):
        for w in self._body.winfo_children():
            w.destroy()

        if not self._students:
            muted(self._body, "No students found.").pack(pady=24)
            return

        for i, s in enumerate(self._students):
            bg = SURFACE if i % 2 == 0 else "#FAFAFA"
            row = ctk.CTkFrame(self._body, fg_color=bg,
                               corner_radius=0, height=38)
            row.pack(fill="x")
            row.pack_propagate(False)

            def _cell(text, width, color=TEXT):
                ctk.CTkLabel(row, text=str(text), font=("", 12),
                             text_color=color, width=width,
                             anchor="w").pack(side="left", padx=(12, 0))

            _cell(s["admission_number"], 130)
            _cell(s["full_name"],        230)
            _cell(s.get("class_name") or "—", 100)
            if self._show_stream_col:
                _cell(s.get("stream") or "—", 80)
            _cell(s.get("gender") or "—", 60)

            actions = ctk.CTkFrame(row, fg_color="transparent")
            actions.pack(side="left", padx=6)

            ctk.CTkButton(actions, text="Edit", width=48, height=26,
                          fg_color="transparent", border_color=BORDER,
                          border_width=1, text_color=TEXT, corner_radius=6,
                          hover_color=BG, font=("", 11),
                          command=lambda st=s: self._open_edit_form(st)
                          ).pack(side="left", padx=(0, 4))

            ctk.CTkButton(actions, text="Transfer", width=68, height=26,
                          fg_color="transparent", border_color=BORDER,
                          border_width=1, text_color=TEXT, corner_radius=6,
                          hover_color=BG, font=("", 11),
                          command=lambda st=s: self._open_transfer(st)
                          ).pack(side="left", padx=(0, 4))

            if s.get("status") == "active":
                ctk.CTkButton(actions, text="Archive", width=64, height=26,
                              fg_color="transparent", border_color=DANGER,
                              border_width=1, text_color=DANGER,
                              corner_radius=6, hover_color="#FEF2F2",
                              font=("", 11),
                              command=lambda st=s: self._do_archive(st)
                              ).pack(side="left")

    # ── Actions ───────────────────────────────────────────────
    def _open_add_form(self):
        StudentForm(self, title="Add student", classes=self._classes,
                    on_save=self._on_save_new)

    def _open_edit_form(self, student):
        StudentForm(self, title="Edit student", classes=self._classes,
                    student=student, on_save=self._on_save_edit)

    def _open_transfer(self, student):
        TransferDialog(self, student=student, classes=self._classes,
                       on_save=self._on_transfer)

    def _on_save_new(self, data):
        ok, msg = create_student(data["full_name"], data["admission_number"],
                                 data["class_id"], data.get("gender"))
        if ok: self._load()
        return ok, msg

    def _on_save_edit(self, data):
        ok, msg = update_student(data["student_id"], data["full_name"],
                                 data["admission_number"], data["class_id"],
                                 data.get("gender"))
        if ok: self._load()
        return ok, msg

    def _on_transfer(self, student_id, class_id):
        transfer_student(student_id, class_id)
        self._load()

    def _do_archive(self, student):
        ConfirmDialog(
            self,
            message=f"Archive {student['full_name']}?\nThey will be hidden from active lists.",
            on_confirm=lambda: [archive_student(student["id"]), self._load()],
        )

    # ── Print / export ────────────────────────────────────────
    def _print_class_list(self):
        PrintExportDialog(self, classes=self._classes)

    # ── Import ────────────────────────────────────────────────
    def _open_import(self):
        ImportDialog(self, classes=self._classes, on_done=self._load)


# ── Student form ──────────────────────────────────────────────
class StudentForm(ctk.CTkToplevel):
    def __init__(self, parent, title, classes, on_save, student=None):
        super().__init__(parent)
        self.title(title)
        self.geometry("460x500")
        self.resizable(False, False)
        self.grab_set()
        self._classes = classes
        self._on_save = on_save
        self._student = student
        self._build()
        if student:
            self._populate(student)

    def _build(self):
        f = ctk.CTkFrame(self, fg_color=BG)
        f.pack(fill="both", expand=True, padx=36, pady=32)

        heading(f, self.title()).pack(anchor="w", pady=(0, 16))

        muted(f, "Full name *").pack(anchor="w")
        self._name = ctk.CTkEntry(f, width=380,
                                   fg_color=SURFACE, border_color=BORDER)
        self._name.pack(anchor="w", pady=(4, 14))

        muted(f, "Admission number *").pack(anchor="w")
        self._adm = ctk.CTkEntry(f, width=380,
                                  fg_color=SURFACE, border_color=BORDER)
        self._adm.pack(anchor="w", pady=(4, 14))

        muted(f, "Class & stream *").pack(anchor="w")
        self._class_labels = [
            f"{c['name']}{' ' + c['stream'] if c.get('stream') else ''}"
            for c in self._classes
        ]
        self._class_var = ctk.StringVar(
            value=self._class_labels[0] if self._class_labels else "")
        ctk.CTkOptionMenu(f, variable=self._class_var,
                          values=self._class_labels, width=380,
                          fg_color=SURFACE, button_color=BORDER,
                          button_hover_color=ACCENT, text_color=TEXT,
                          dropdown_fg_color=SURFACE,
                          ).pack(anchor="w", pady=(4, 14))

        muted(f, "Gender (optional)").pack(anchor="w")
        self._gender_var = ctk.StringVar(value="—")
        ctk.CTkOptionMenu(f, variable=self._gender_var,
                          values=["—", "M", "F"], width=380,
                          fg_color=SURFACE, button_color=BORDER,
                          button_hover_color=ACCENT, text_color=TEXT,
                          dropdown_fg_color=SURFACE,
                          ).pack(anchor="w", pady=(4, 14))

        self._error = ctk.CTkLabel(f, text="",
                                    text_color=DANGER, font=("", 12))
        self._error.pack(anchor="w")

        btn_row = ctk.CTkFrame(f, fg_color="transparent")
        btn_row.pack(fill="x", pady=(12, 0))
        ghost_btn(btn_row, "Cancel", command=self.destroy, width=100
                  ).pack(side="left")
        primary_btn(btn_row, "Save", command=self._submit, width=100
                    ).pack(side="right")

    def _populate(self, s):
        self._name.insert(0, s["full_name"])
        self._adm.insert(0, s["admission_number"])
        if s.get("gender"):
            self._gender_var.set(s["gender"])
        cls_label = f"{s.get('class_name', '')}{' ' + s.get('stream', '') if s.get('stream') else ''}".strip()
        if cls_label in self._class_labels:
            self._class_var.set(cls_label)

    def _submit(self):
        self._error.configure(text="")
        name   = self._name.get().strip()
        adm    = self._adm.get().strip()
        gender = self._gender_var.get()
        gender = None if gender == "—" else gender
        cls_label = self._class_var.get()
        match = next((c for c in self._classes
                      if f"{c['name']}{' ' + c['stream'] if c.get('stream') else ''}" == cls_label),
                     None)
        if not match:
            self._error.configure(text="Please select a valid class.")
            return
        data = {"full_name": name, "admission_number": adm,
                "class_id": match["id"], "gender": gender}
        if self._student:
            data["student_id"] = self._student["id"]
        ok, msg = self._on_save(data)
        if ok:
            self.destroy()
        else:
            self._error.configure(text=msg)


# ── Transfer dialog ───────────────────────────────────────────
class TransferDialog(ctk.CTkToplevel):
    def __init__(self, parent, student, classes, on_save):
        super().__init__(parent)
        self.title("Transfer student")
        self.geometry("460x260")
        self.resizable(False, False)
        self.grab_set()
        self._student = student
        self._classes = classes
        self._on_save = on_save
        self._build()

    def _build(self):
        f = ctk.CTkFrame(self, fg_color=BG)
        f.pack(fill="both", expand=True, padx=36, pady=32)

        heading(f, "Transfer student", size=16).pack(anchor="w", pady=(0, 4))
        muted(f, f"Moving: {self._student['full_name']}").pack(
            anchor="w", pady=(0, 14))

        muted(f, "New class & stream").pack(anchor="w")
        self._class_labels = [
            f"{c['name']}{' ' + c['stream'] if c.get('stream') else ''}"
            for c in self._classes
        ]
        self._class_var = ctk.StringVar(
            value=self._class_labels[0] if self._class_labels else "")
        ctk.CTkOptionMenu(f, variable=self._class_var,
                          values=self._class_labels, width=380,
                          fg_color=SURFACE, button_color=BORDER,
                          button_hover_color=ACCENT, text_color=TEXT,
                          dropdown_fg_color=SURFACE,
                          ).pack(anchor="w", pady=(4, 16))

        btn_row = ctk.CTkFrame(f, fg_color="transparent")
        btn_row.pack(fill="x")
        ghost_btn(btn_row, "Cancel", command=self.destroy, width=100
                  ).pack(side="left")
        primary_btn(btn_row, "Transfer", command=self._submit, width=100
                    ).pack(side="right")

    def _submit(self):
        cls_label = self._class_var.get()
        match = next((c for c in self._classes
                      if f"{c['name']}{' ' + c['stream'] if c.get('stream') else ''}" == cls_label),
                     None)
        if match:
            self._on_save(self._student["id"], match["id"])
            self.destroy()


# ── Print / export dialog ─────────────────────────────────────
class PrintExportDialog(ctk.CTkToplevel):
    def __init__(self, parent, classes):
        super().__init__(parent)
        self.title("Print / export class list")
        self.geometry("520x500")
        self.resizable(False, False)
        self.grab_set()
        self._classes = classes
        self._build()

    def _build(self):
        f = ctk.CTkFrame(self, fg_color=BG)
        f.pack(fill="both", expand=True, padx=36, pady=32)

        heading(f, "Print / export", size=16).pack(anchor="w", pady=(0, 14))

        # Class selector
        seen, self._class_names = [], []
        for c in self._classes:
            if c["name"] not in seen:
                seen.append(c["name"])
                self._class_names.append(c["name"])

        muted(f, "Class").pack(anchor="w")
        options = ["All classes"]
        for name in self._class_names:
            streams = [c for c in self._classes
                       if c["name"] == name and c.get("stream")]
            if streams:
                options.append(f"{name} — all streams")
                for c in streams:
                    options.append(f"{name} {c['stream']}")
            else:
                options.append(name)
        self._class_var = ctk.StringVar(value="All classes")
        ctk.CTkOptionMenu(f, variable=self._class_var,
                          values=options, width=430,
                          fg_color=SURFACE, button_color=BORDER,
                          text_color=TEXT, dropdown_fg_color=SURFACE,
                          ).pack(anchor="w", pady=(4, 14))

        # Gender filter
        muted(f, "Gender").pack(anchor="w")
        self._gender_var = ctk.StringVar(value="All")
        ctk.CTkOptionMenu(f, variable=self._gender_var,
                          values=["All", "Male only", "Female only"],
                          width=430, fg_color=SURFACE,
                          button_color=BORDER, text_color=TEXT,
                          dropdown_fg_color=SURFACE,
                          ).pack(anchor="w", pady=(4, 14))

        # Export format
        muted(f, "Action").pack(anchor="w")
        self._format_var = ctk.StringVar(value="Save as PDF")
        ctk.CTkOptionMenu(f, variable=self._format_var,
                          values=["Save as PDF", "Save as Excel (.xlsx)",
                                  "Save as Word (.docx)", "Print to printer"],
                          width=430, fg_color=SURFACE,
                          button_color=BORDER, text_color=TEXT,
                          dropdown_fg_color=SURFACE,
                          ).pack(anchor="w", pady=(4, 16))

        self._status = ctk.CTkLabel(f, text="", font=("", 12),
                                     text_color=TEXT_MUTED)
        self._status.pack(anchor="w", pady=(0, 8))

        btn_row = ctk.CTkFrame(f, fg_color="transparent")
        btn_row.pack(fill="x")
        ghost_btn(btn_row, "Cancel", command=self.destroy,
                  width=100).pack(side="left")
        primary_btn(btn_row, "Go", command=self._go,
                    width=100).pack(side="right")

    def _resolve_class(self):
        sel = self._class_var.get().strip()
        class_id = None
        class_name_filter = None
        if sel == "All classes":
            pass
        elif "— all streams" in sel:
            class_name_filter = sel.replace(" — all streams", "").strip()
        else:
            match = next(
                (c for c in self._classes
                 if (c.get("stream") and f"{c['name']} {c['stream']}" == sel)
                 or (not c.get("stream") and c["name"] == sel)),
                None)
            if match:
                class_id = match["id"]
        return class_id, class_name_filter

    def _get_students(self, class_id, class_name_filter):
        from routes.students import get_students as gs
        if class_id:
            students = gs(class_id=class_id)
        elif class_name_filter:
            students = [s for s in gs() if s.get("class_name") == class_name_filter]
        else:
            students = gs()
        # Gender filter
        g = self._gender_var.get()
        if g == "Male only":
            students = [s for s in students if s.get("gender") == "M"]
        elif g == "Female only":
            students = [s for s in students if s.get("gender") == "F"]
        return students

    def _go(self):
        fmt = self._format_var.get()
        class_id, class_name_filter = self._resolve_class()

        if fmt == "Save as PDF":
            self._export_pdf(class_id, class_name_filter)
        elif fmt == "Save as Excel (.xlsx)":
            self._export_excel(class_id, class_name_filter)
        elif fmt == "Save as Word (.docx)":
            self._export_docx(class_id, class_name_filter)
        elif fmt == "Print to printer":
            self._print_direct(class_id, class_name_filter)

    def _export_pdf(self, class_id, class_name_filter):
        sel = self._class_var.get().strip().replace(" ", "_").replace("—", "")
        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
            initialfile=f"classlist_{sel}.pdf",
            title="Save PDF as",
        )
        if not path:
            return
        self._status.configure(text="Generating PDF…")
        self.update()
        try:
            # Pass gender filter to PDF generator
            gender = self._gender_var.get()
            gender_filter = None
            if gender == "Male only":   gender_filter = "M"
            elif gender == "Female only": gender_filter = "F"
            generate_class_list(path, class_id=class_id,
                                 class_name_filter=class_name_filter,
                                 gender_filter=gender_filter)
            self._status.configure(text=f"✓ Saved: {os.path.basename(path)}",
                                   text_color=SUCCESS)
            os.system(f'open "{path}"' if os.name != "nt"
                      else f'start "" "{path}"')
            self.after(1800, self.destroy)
        except Exception as e:
            self._status.configure(text=f"Error: {e}", text_color=DANGER)

    def _export_excel(self, class_id, class_name_filter):
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        students = self._get_students(class_id, class_name_filter)
        if not students:
            self._status.configure(text="No students found.", text_color=DANGER)
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")],
            initialfile="classlist.xlsx", title="Save Excel as")
        if not path:
            return
        wb  = openpyxl.Workbook()
        ws  = wb.active
        ws.title = "Class List"
        from routes.settings import get_setting
        from routes.terms import get_current_term
        school = get_setting("school_name", "GradeVault")
        term   = get_current_term()
        ws.append([school])
        ws["A1"].font = Font(bold=True, size=14)
        ws.append([f"Class List — {self._class_var.get()}"])
        if term:
            ws.append([f"Term {term['term']}, {term['year']}  ·  {len(students)} student(s)"])
        ws.append([])
        headers = ["Adm. No.", "Full Name", "Class", "Stream", "Gender"]
        ws.append(headers)
        hdr_fill = PatternFill("solid", fgColor="EEF2FF")
        for cell in ws[ws.max_row]:
            cell.font = Font(bold=True, color="4F46E5")
            cell.fill = hdr_fill
        for s in students:
            ws.append([
                s["admission_number"], s["full_name"],
                s.get("class_name") or "", s.get("stream") or "",
                s.get("gender") or "",
            ])
        ws.column_dimensions["A"].width = 18
        ws.column_dimensions["B"].width = 28
        ws.column_dimensions["C"].width = 12
        ws.column_dimensions["D"].width = 10
        ws.column_dimensions["E"].width = 10
        wb.save(path)
        self._status.configure(text=f"✓ Saved: {os.path.basename(path)}",
                               text_color=SUCCESS)
        os.system(f'open "{path}"' if os.name != "nt" else f'start "" "{path}"')
        self.after(1800, self.destroy)

    def _export_docx(self, class_id, class_name_filter):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            self._status.configure(
                text="python-docx not installed. Run: pip install python-docx",
                text_color=DANGER)
            return
        students = self._get_students(class_id, class_name_filter)
        if not students:
            self._status.configure(text="No students found.", text_color=DANGER)
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            initialfile="classlist.docx", title="Save Word doc as")
        if not path:
            return
        from routes.settings import get_setting
        from routes.terms import get_current_term
        doc    = Document()
        school = get_setting("school_name", "GradeVault")
        term   = get_current_term()
        title  = doc.add_heading(school, level=1)
        title.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        doc.add_heading(f"Class List — {self._class_var.get()}", level=2)
        if term:
            doc.add_paragraph(f"Term {term['term']}, {term['year']}  ·  {len(students)} student(s)")
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Adm. No.", "Full Name", "Class", "Stream", "Gender"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
        for s in students:
            row = table.add_row().cells
            row[0].text = s["admission_number"]
            row[1].text = s["full_name"]
            row[2].text = s.get("class_name") or ""
            row[3].text = s.get("stream") or ""
            row[4].text = s.get("gender") or ""
        doc.save(path)
        self._status.configure(text=f"✓ Saved: {os.path.basename(path)}",
                               text_color=SUCCESS)
        os.system(f'open "{path}"' if os.name != "nt" else f'start "" "{path}"')
        self.after(1800, self.destroy)

    def _print_direct(self, class_id, class_name_filter):
        import tempfile
        self._status.configure(text="Preparing print job…")
        self.update()
        try:
            gender = self._gender_var.get()
            gender_filter = None
            if gender == "Male only":    gender_filter = "M"
            elif gender == "Female only": gender_filter = "F"
            with tempfile.NamedTemporaryFile(suffix=".pdf",
                                             delete=False) as tmp:
                tmp_path = tmp.name
            generate_class_list(tmp_path, class_id=class_id,
                                 class_name_filter=class_name_filter,
                                 gender_filter=gender_filter)
            if os.name == "nt":
                os.startfile(tmp_path, "print")
            else:
                os.system(f'lpr "{tmp_path}"')
            self._status.configure(text="✓ Sent to printer.",
                                   text_color=SUCCESS)
            self.after(1800, self.destroy)
        except Exception as e:
            self._status.configure(text=f"Error: {e}", text_color=DANGER)


# ── Import dialog ─────────────────────────────────────────────
class ImportDialog(ctk.CTkToplevel):
    def __init__(self, parent, classes, on_done):
        super().__init__(parent)
        self.title("Import students from CSV / Excel")
        self.geometry("560x520")
        self.resizable(False, False)
        self.grab_set()
        self._classes = classes
        self._on_done = on_done
        self._rows    = []
        self._build()

    def _build(self):
        f = ctk.CTkFrame(self, fg_color=BG)
        f.pack(fill="both", expand=True, padx=36, pady=32)

        heading(f, "Import students", size=16).pack(anchor="w", pady=(0, 4))
        muted(f, "Upload a CSV or Excel file with student data.").pack(
            anchor="w", pady=(0, 6))

        hint = ctk.CTkFrame(f, fg_color=ACCENT_BG, corner_radius=6)
        hint.pack(fill="x", pady=(0, 12))
        ctk.CTkLabel(hint,
                     text="Required columns:  Full Name  |  Admission number  |  Gender (optional)",
                     font=("", 11), text_color=ACCENT,
                     ).pack(padx=10, pady=6, anchor="w")

        ghost_btn(f, "Download template CSV",
                  command=self._download_template, width=180
                  ).pack(anchor="w", pady=(0, 12))

        muted(f, "Assign to class:").pack(anchor="w")
        class_labels = [
            f"{c['name']}{' ' + c['stream'] if c.get('stream') else ''}"
            for c in self._classes
        ]
        self._class_var = ctk.StringVar(
            value=class_labels[0] if class_labels else "")
        ctk.CTkOptionMenu(f, variable=self._class_var,
                          values=class_labels if class_labels else ["—"],
                          width=430, fg_color=SURFACE, button_color=BORDER,
                          text_color=TEXT, dropdown_fg_color=SURFACE,
                          ).pack(anchor="w", pady=(4, 12))

        file_row = ctk.CTkFrame(f, fg_color="transparent")
        file_row.pack(fill="x", pady=(0, 8))
        self._file_label = muted(file_row, "No file selected")
        self._file_label.pack(side="left", expand=True, anchor="w")
        ghost_btn(file_row, "Choose file",
                  command=self._pick_file, width=110).pack(side="right")

        self._preview = ctk.CTkTextbox(f, width=460, height=100,
                                        fg_color=SURFACE, border_color=BORDER,
                                        font=("", 11), state="disabled")
        self._preview.pack(pady=(0, 8))

        self._status = ctk.CTkLabel(f, text="", font=("", 12),
                                     text_color=TEXT_MUTED)
        self._status.pack(anchor="w", pady=(0, 8))

        btn_row = ctk.CTkFrame(f, fg_color="transparent")
        btn_row.pack(fill="x")
        ghost_btn(btn_row, "Cancel", command=self.destroy, width=100
                  ).pack(side="left")
        self._import_btn = primary_btn(btn_row, "Import students",
                                        command=self._do_import, width=140)
        self._import_btn.pack(side="right")
        self._import_btn.configure(state="disabled")

    def _download_template(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="students_template.csv",
            title="Save template as")
        if path:
            with open(path, "w") as fh:
                fh.write(sample_csv_template())
            self._status.configure(
                text=f"Template saved.", text_color=SUCCESS)

    def _pick_file(self):
        path = filedialog.askopenfilename(
            filetypes=[("CSV / Excel", "*.csv *.xlsx *.xls")],
            title="Select student file")
        if not path:
            return
        self._file_label.configure(text=os.path.basename(path))
        rows, err = read_students_from_file(path)
        self._rows = rows
        self._preview.configure(state="normal")
        self._preview.delete("1.0", "end")
        if rows:
            lines = [f"{r['admission_number']}  |  {r['full_name']}  |  {r.get('gender') or '—'}"
                     for r in rows[:8]]
            if len(rows) > 8:
                lines.append(f"... and {len(rows) - 8} more")
            self._preview.insert("1.0", "\n".join(lines))
            self._import_btn.configure(state="normal")
            msg = f"{len(rows)} student(s) ready to import."
            if err:
                msg += f"\nWarnings:\n{err}"
            self._status.configure(text=msg, text_color=TEXT_MUTED)
        else:
            self._preview.insert("1.0", err or "No valid rows found.")
            self._import_btn.configure(state="disabled")
            self._status.configure(text="Fix the file and try again.",
                                   text_color=DANGER)
        self._preview.configure(state="disabled")

    def _do_import(self):
        if not self._rows:
            return
        cls_label = self._class_var.get()
        cls = next((c for c in self._classes
                    if f"{c['name']}{' ' + c['stream'] if c.get('stream') else ''}" == cls_label),
                   None)
        if not cls:
            self._status.configure(text="Select a valid class first.",
                                   text_color=DANGER)
            return
        success, skipped = 0, 0
        for r in self._rows:
            ok, _ = create_student(r["full_name"], r["admission_number"],
                                   cls["id"], r.get("gender"))
            if ok: success += 1
            else:  skipped += 1
        msg = f"✓ Imported {success} student(s)."
        if skipped:
            msg += f"  {skipped} skipped (duplicate admission numbers)."
        self._status.configure(text=msg, text_color=SUCCESS)
        self._import_btn.configure(state="disabled")
        self.after(1800, lambda: [self.destroy(), self._on_done()])


# ── Confirm dialog ────────────────────────────────────────────
class ConfirmDialog(ctk.CTkToplevel):
    def __init__(self, parent, message, on_confirm):
        super().__init__(parent)
        self.title("Confirm")
        self.geometry("440x200")
        self.resizable(False, False)
        self.grab_set()
        f = ctk.CTkFrame(self, fg_color=BG)
        f.pack(fill="both", expand=True, padx=36, pady=32)
        label(f, message, size=13).pack(anchor="w", pady=(0, 20))
        btn_row = ctk.CTkFrame(f, fg_color="transparent")
        btn_row.pack(fill="x")
        ghost_btn(btn_row, "Cancel", command=self.destroy,
                  width=100).pack(side="left")
        ctk.CTkButton(btn_row, text="Confirm", width=100, height=38,
                      fg_color=DANGER, hover_color="#DC2626",
                      corner_radius=8,
                      command=lambda: [on_confirm(), self.destroy()]
                      ).pack(side="right")
